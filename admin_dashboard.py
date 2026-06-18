import os
import io
import re
import sys
import logging
import json
import requests
import urllib3
from datetime import datetime
from functools import wraps
from flask import (Flask, request, redirect, url_for, session, jsonify,
                   render_template_string, flash, send_file)
from sqlalchemy import or_
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from database import get_db
from models import OnboardingSession, Question, Base

# Load environment variables
load_dotenv()
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Configuration
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "spike2025")
SECRET_KEY = os.getenv("SECRET_KEY", "spike-admin-secret-key-change-in-prod")
SYSTEM_TYPES = ["Generic", "AD-Azure", "Target DB", "SAP", "LDAP"]

# Azure OpenAI Configuration
AZURE_LLM_URL = os.getenv("AZURE_LLM_URL", "https://spikeiam-genai-resource.cognitiveservices.azure.com/openai/responses?api-version=2025-04-01-preview")
AZURE_API_KEY = os.getenv("AZURE_API_KEY", "")
AZURE_LLM_MODEL = os.getenv("AZURE_LLM_MODEL", "gpt-5.4-mini")

# Flask app
app = Flask(__name__)
app.secret_key = SECRET_KEY

# =============================================================================
# LOGGING
# =============================================================================

LOG_LEVEL = os.getenv("LOG_LEVEL", "DEBUG").upper()

logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.DEBUG),
    format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
    stream=sys.stdout,
    force=True,
)

logger = logging.getLogger("spike-admin")
app.logger.handlers = logger.handlers
app.logger.setLevel(getattr(logging, LOG_LEVEL, logging.DEBUG))

logger.info("Spike Admin dashboard module loaded")
logger.info("Configured system types: %s", SYSTEM_TYPES)
logger.info("Azure model configured as: %s", AZURE_LLM_MODEL)
logger.info("Azure API key present: %s", "yes" if bool(AZURE_API_KEY) else "no")

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('authenticated'):
            logger.debug("Unauthenticated request redirected to login: %s %s", request.method, request.path)
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


@app.before_request
def log_request_start():
    logger.info("REQ  %s %s from %s", request.method, request.path, request.remote_addr)


@app.after_request
def log_request_end(response):
    logger.info("RES  %s %s -> %s", request.method, request.path, response.status_code)
    return response


# =============================================================================
# SOLUTION DESIGN TEMPLATE STRUCTURE
# =============================================================================

SOLUTION_DESIGN_TEMPLATE_PROMPT = """
You MUST generate the Solution Design document following the EFG Solution Design structure.

Critical grounding rules:
- Use ONLY the interview data provided in the current OnboardingSession.
- Do NOT fill a field with generic assumptions.
- Before writing [TO BE DEFINED], verify that no interview answer directly or indirectly provides that information.
- If an answer partially covers a section, write the available part and use [TO BE DEFINED] only for the missing sub-fields.
- Never mark an entire section as [TO BE DEFINED] if at least one interview answer can be used inside that section.
- Keep the wording professional and consistent with enterprise IAM / IGA documentation.
- Output valid markdown only. Do not include code fences.

DOCUMENT STRUCTURE TO FOLLOW:

# Purpose of the Document

{company} has started a broad project initiative with the aim of installing and integrating a new IGA solution within the company context aimed at managing the life cycle of users, organizations and access rights in a more effective and efficient way, as well as the internal processes related to them, provisioning and synchronization from the main target systems, RBAC model and SoD rules implementation.

The purpose of the document, after a brief introduction on the project requirements, is to describe in depth the solution adopted both in terms of available functionalities and in architectural terms, describing at a high level the individual components and their interaction and integration mechanisms. In the last part, it also describes the chosen approach in terms of integration of the target systems with a detailed explanation of the configured components.

The document is intended for an audience with technical knowledge in the IT and UAM areas of {company}.

# References

| DOCUMENT TITLE | DESCRIPTION |
|---|---|
| {company} - IGA - Installation Manual | Installation Manual of the system. |
| {company} - IGA - Requirements Specifics | Descriptive document of as-is context, functional and non-functional requirements, constraints and project deliverables. |

# Definitions, Abbreviations and Acronyms

| TERM / ABBREVIATION | DEFINITION |
|---|---|
| ABAC | Attribute-Based Access Control |
| AD | Active Directory |
| Azure AD / Entra ID | Microsoft cloud identity platform used for authentication, synchronization, or license management where applicable |
| DB | Database |
| IGA | Identity Governance and Administration |
| RBAC | Role-Based Access Control |
| SoD | Segregation of Duties |
| UAM | User Access Management |

# Project Goals

## Current State (As-Is)
Describe the current state using the interview data:
- Current account creation, modification, disabling, and access assignment process
- Current manual, semi-automatic, or automatic steps
- Available authoritative sources only if mentioned
- Current target-system management model
- Current identity/account linking logic, if available

## Current Issues
Describe issues supported by interview evidence. If no issue is explicit, write:
[TO BE DEFINED]

Possible issue categories to populate only when supported by the interview:
- User productivity lost due to manual provisioning process
- Unnecessary administrative overhead
- Reduced security or audit risk
- Poor user experience
- Lack of integration or inconsistent lifecycle process

## Project Objectives
Describe objectives that logically follow from the interview data:
- Life cycle management of identities and accounts
- Management of target systems, accounts, and access rights
- Authorization management and RBAC/ABAC model
- Governance and recertification campaigns
- SoD rules
- Reporting, auditing, and monitoring activities

# Solution Description

## One Identity Manager Overview
Describe the IGA capabilities relevant to the received interview:
- Identity and organizational unit life cycle
- Access rights management
- Access requests and approval workflows
- Management of rules and roles (RBAC/ABAC)
- Target systems management
- Recertification of access rights
- SoD rules
- Auditing, reporting, and monitoring

## One Identity Manager Architecture
Describe the standard One Identity Manager architecture:
- Central Repository
- Web Server and Application Server
- Job Server / Provisioning Server
- Object Layer
- Administrative Clients

# Architecture Overview

## Project Environments
Use interview answers about environments, e.g. PROD, DEV, TEST, UAT, sandbox, or additional instances. If only partial information is provided, list known environments and mark missing details as [TO BE DEFINED].

## Test Environment Details
For each component, fill only information available from the interview:
- Workstation: Processor, Memory, Hard drive, OS, Additional Software
- Job Server: Processor, Memory, Hard drive, OS, Additional Software
- Application Server: Processor, Memory, Hard drive, OS, Additional Software
- Database: Processor, Memory, Hard drive, OS, Additional Software

## Production Environment Details
Same structure as Test Environment. Fill only information available from the interview.

## High Availability, Backup and Disaster Recovery
Use interview evidence if present. Otherwise write [TO BE DEFINED].

# Integration with Authoritative Sources and Target Systems

## Integration with Authoritative Sources
Describe authoritative sources only if mentioned in the interview. For each source:
- Source name
- Connector type
- Connection details
- Read/write mode
- Synchronization schedule
- Identity linking logic

If no authoritative source is mentioned, write [TO BE DEFINED].

## Integration with Target Systems

### Target System Summary
- Company: {company}
- Target System name: {target_system}
- System type: {system_type}
- Deployment model
- Network location
- Authentication mechanism
- Authorization/profiling mechanism
- Available environments or instances

### Connector and Connection Configuration
Based on system type and interview data:
- For AD-Azure: AD / Azure AD connector details, domains, domain controllers if known, Azure AD Connect schedule, synchronization scope, relevant network/firewall constraints.
- For LDAP: LDAP connector details, domains, synchronization scope, relevant network/firewall constraints.
- For Target DB: Generic Database Connector details, database technology, server/database/schema/table names if available, service account, expected privileges, read/write mode.
- For SAP: SAP connector details if available.
- For Generic: connector type must be derived from interview answers; otherwise [TO BE DEFINED].

### Account Definition and Identity Linking
Use interview answers about:
- Unique username/key
- SAMAccountName / UPN convention
- Employee ID, Central Account, login name, or application key
- Multiple account/profile behavior
- Account definition name, if inferable from target system name

### Attribute Mapping Logic
Use interview answers about:
- Mandatory and optional attributes
- User status attributes and values
- OU placement rules
- Group assignment rules
- License assignment rules
- Synchronization scope
- Database user/profile/access-right tables

### Access Rights / Entitlements
Use interview answers about:
- Groups, roles, profiles, permissions, access-right tables
- Naming convention
- Description availability
- Business/application owner if mentioned
- Standard/default assignments
- Exceptions or special cases

### Synchronization and Provisioning
Use interview answers about:
- Synchronization schedule
- Read-only or read/write behavior
- Manual, semi-automatic, or automatic provisioning
- Service account privileges
- Constraints and custom scenarios

# RBAC Model

## Import of Business Roles
Use interview data about groups, roles, access rights, naming convention, default assignment rules, and entitlement structures. If no business-role import model is described, write [TO BE DEFINED].

## Managing Business Roles via Web Portal
Describe request, approval, assignment, and removal workflows only if supported by the interview. Otherwise write [TO BE DEFINED].

# Joiner, Mover and Leaver Processes

## Account Automation
For AD-Azure, describe AD account automation using OU, group, license, SAMAccountName/UPN, and lifecycle answers.
For Target DB, describe database user-record creation, modification, status management, and access-right assignment answers.
For any missing sub-field, use [TO BE DEFINED].

## Management of Admin Accounts
Use admin, privileged profile, technical account, service account, or application support answers. Otherwise [TO BE DEFINED].

## Movers - Access Revocation
Use answers about department changes, role changes, OU moves, profile updates, and revocation. Otherwise [TO BE DEFINED].

## Identity and Account Deactivation
Use answers about disabling, termination, inactive status, disabled OU, assignment end date, account lock, and deletion. Otherwise [TO BE DEFINED].

## Automatic Deactivation and Deletion
Use answers about retention periods, deletion timelines, status values, and automatic processes. Otherwise [TO BE DEFINED].
"""

# =============================================================================
# TEMPLATES
# =============================================================================

BASE_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ page_title }} - Spike Admin</title>
    <style>
        :root {
            --bg-primary: #0f1117;
            --bg-secondary: #1a1d27;
            --bg-card: #1e2130;
            --bg-hover: #252838;
            --border-color: #2a2d3a;
            --text-primary: #e4e4e7;
            --text-secondary: #a1a1aa;
            --text-muted: #71717a;
            --accent-primary: #6366f1;
            --accent-hover: #4f46e5;
            --accent-success: #10b981;
            --accent-warning: #f59e0b;
            --accent-danger: #ef4444;
            --sidebar-width: 260px;
        }
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: var(--bg-primary);
            color: var(--text-primary);
            display: flex;
            min-height: 100vh;
        }
        .sidebar {
            width: var(--sidebar-width);
            background: var(--bg-secondary);
            border-right: 1px solid var(--border-color);
            padding: 24px 16px;
            position: fixed;
            top: 0;
            left: 0;
            height: 100vh;
            overflow-y: auto;
            display: flex;
            flex-direction: column;
        }
        .sidebar-brand {
            font-size: 1.25rem;
            font-weight: 700;
            color: var(--accent-primary);
            padding: 0 12px 24px;
            border-bottom: 1px solid var(--border-color);
            margin-bottom: 24px;
        }
        .nav-item {
            display: flex;
            align-items: center;
            gap: 12px;
            padding: 10px 12px;
            border-radius: 8px;
            color: var(--text-secondary);
            text-decoration: none;
            font-size: 0.9rem;
            margin-bottom: 4px;
            transition: all 0.15s ease;
        }
        .nav-item:hover { background: var(--bg-hover); color: var(--text-primary); }
        .nav-item.active { background: var(--accent-primary); color: white; }
        .nav-icon { display: inline-flex; align-items: center; }
        .nav-icon svg { width: 18px; height: 18px; }
        .sidebar-footer {
            margin-top: auto;
            padding-top: 24px;
            border-top: 1px solid var(--border-color);
        }
        .main-content {
            margin-left: var(--sidebar-width);
            flex: 1;
            padding: 32px;
            max-width: calc(100vw - var(--sidebar-width));
        }
        .page-header {
            margin-bottom: 32px;
        }
        .page-header h1 {
            font-size: 1.75rem;
            font-weight: 700;
            color: var(--text-primary);
        }
        .page-header p {
            color: var(--text-muted);
            margin-top: 4px;
        }
        .card {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 24px;
            margin-bottom: 24px;
        }
        .card-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
        }
        .card-header h2, .card-header h3 {
            font-size: 1.1rem;
            font-weight: 600;
        }
        .btn {
            padding: 8px 16px;
            border-radius: 8px;
            border: none;
            font-size: 0.85rem;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.15s ease;
            text-decoration: none;
            display: inline-flex;
            align-items: center;
            gap: 6px;
        }
        .btn-primary { background: var(--accent-primary); color: white; }
        .btn-primary:hover { background: var(--accent-hover); }
        .btn-danger { background: var(--accent-danger); color: white; }
        .btn-danger:hover { background: #dc2626; }
        .btn-sm { padding: 6px 12px; font-size: 0.8rem; }
        .table-container { overflow-x: auto; }
        table {
            width: 100%;
            border-collapse: collapse;
        }
        th, td {
            padding: 12px 16px;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }
        th {
            font-size: 0.8rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: var(--text-muted);
        }
        td { font-size: 0.9rem; color: var(--text-secondary); }
        .badge {
            padding: 4px 10px;
            border-radius: 12px;
            font-size: 0.75rem;
            font-weight: 500;
        }
        .badge-generic { background: #1e3a5f; color: #60a5fa; }
        .badge-ad-azure { background: #1e3a2a; color: #34d399; }
        .badge-target-db { background: #3b2a1e; color: #fbbf24; }
        .badge-sap { background: #2a1e3b; color: #a78bfa; }
        .badge-ldap {background: rgba(43, 25, 102, 0.15); color: #1ab6ef;}
        .badge-active { background: #064e3b; color: #6ee7b7; }
        .badge-completed { background: #1e3a5f; color: #60a5fa; }
        .badge-abandoned { background: #3b1e1e; color: #fca5a5; }
        .stats-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 16px;
            margin-bottom: 24px;
        }
        .stat-card {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 20px;
        }
        .stat-value {
            font-size: 2rem;
            font-weight: 700;
            color: var(--text-primary);
        }
        .stat-label {
            font-size: 0.85rem;
            color: var(--text-muted);
            margin-top: 4px;
        }
        .form-group { margin-bottom: 16px; }
        .form-group label {
            display: block;
            font-size: 0.85rem;
            font-weight: 500;
            color: var(--text-secondary);
            margin-bottom: 6px;
        }
        .form-control {
            width: 100%;
            padding: 10px 14px;
            border-radius: 8px;
            border: 1px solid var(--border-color);
            background: var(--bg-primary);
            color: var(--text-primary);
            font-size: 0.9rem;
        }
        .form-control:focus {
            outline: none;
            border-color: var(--accent-primary);
        }
        select.form-control { appearance: none; }
        textarea.form-control { min-height: 80px; resize: vertical; }
        .modal-overlay {
            display: none;
            position: fixed;
            top: 0; left: 0; right: 0; bottom: 0;
            background: rgba(0,0,0,0.7);
            z-index: 1000;
            align-items: center;
            justify-content: center;
        }
        .modal-overlay.active { display: flex; }
        .modal {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 32px;
            width: 90%;
            max-width: 600px;
            max-height: 80vh;
            overflow-y: auto;
        }
        .modal h2 { margin-bottom: 20px; }
        .toast-container {
            position: fixed;
            top: 24px;
            right: 24px;
            z-index: 2000;
        }
        .toast {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 8px;
            padding: 12px 20px;
            margin-bottom: 8px;
            font-size: 0.85rem;
            animation: slideIn 0.3s ease;
        }
        .toast.success { border-left: 4px solid var(--accent-success); }
        .toast.error { border-left: 4px solid var(--accent-danger); }
        @keyframes slideIn {
            from { transform: translateX(100%); opacity: 0; }
            to { transform: translateX(0); opacity: 1; }
        }
        .empty-state {
            text-align: center;
            padding: 48px 24px;
            color: var(--text-muted);
        }
        .empty-icon { font-size: 3rem; margin-bottom: 16px; }
    </style>
</head>
<body>
    <nav class="sidebar">
        <div class="sidebar-brand">Spike Admin</div>
        <a href="{{ url_for('dashboard') }}" class="nav-item {{ 'active' if active_page == 'dashboard' else '' }}">
            <span class="nav-icon"><svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke-width="1.5" stroke="currentColor"><path stroke-linecap="round" stroke-linejoin="round" d="M3 13.125C3 12.504 3.504 12 4.125 12h2.25c.621 0 1.125.504 1.125 1.125v6.75C7.5 20.496 6.996 21 6.375 21h-2.25A1.125 1.125 0 013 19.875v-6.75zM9.75 8.625c0-.621.504-1.125 1.125-1.125h2.25c.621 0 1.125.504 1.125 1.125v11.25c0 .621-.504 1.125-1.125 1.125h-2.25a1.125 1.125 0 01-1.125-1.125V8.625zM16.5 4.125c0-.621.504-1.125 1.125-1.125h2.25C20.496 3 21 3.504 21 4.125v15.75c0 .621-.504 1.125-1.125 1.125h-2.25a1.125 1.125 0 01-1.125-1.125V4.125z"/></svg></span>
            <span>Overview</span>
        </a>
        <a href="{{ url_for('questions_page') }}" class="nav-item {{ 'active' if active_page == 'questions' else '' }}">
            <span class="nav-icon"><svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke-width="1.5" stroke="currentColor"><path stroke-linecap="round" stroke-linejoin="round" d="M9.879 7.519c1.171-1.025 3.071-1.025 4.242 0 1.172 1.025 1.172 2.687 0 3.712-.203.179-.43.326-.67.442-.745.361-1.45.999-1.45 1.827v.75M21 12a9 9 0 11-18 0 9 9 0 0118 0zm-9 5.25h.008v.008H12v-.008z"/></svg></span>
            <span>Questions</span>
        </a>
        <a href="{{ url_for('interviews_page') }}" class="nav-item {{ 'active' if active_page == 'interviews' else '' }}">
            <span class="nav-icon"><svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke-width="1.5" stroke="currentColor"><path stroke-linecap="round" stroke-linejoin="round" d="M20.25 8.511c.884.284 1.5 1.128 1.5 2.097v4.286c0 1.136-.847 2.1-1.98 2.193-.34.027-.68.052-1.02.072v3.091l-3-3c-1.354 0-2.694-.055-4.02-.163a2.115 2.115 0 01-.825-.242m9.345-8.334a2.126 2.126 0 00-.476-.095 48.64 48.64 0 00-8.048 0c-1.131.094-1.976 1.057-1.976 2.192v4.286c0 .837.46 1.58 1.155 1.951m9.345-8.334V6.637c0-1.621-1.152-3.026-2.76-3.235A48.455 48.455 0 0011.25 3c-2.115 0-4.198.137-6.24.402-1.608.209-2.76 1.614-2.76 3.235v6.226c0 1.621 1.152 3.026 2.76 3.235.577.075 1.157.14 1.74.194V21l4.155-4.155"/></svg></span>
            <span>Interviews</span>
        </a>
        <a href="{{ url_for('clients_page') }}" class="nav-item {{ 'active' if active_page == 'clients' else '' }}">
            <span class="nav-icon"><svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke-width="1.5" stroke="currentColor"><path stroke-linecap="round" stroke-linejoin="round" d="M2.25 21h19.5M3.75 3v18m4.5-18v18m4.5-18v18m4.5-18v18m4.5-18v18M3.75 3h16.5M3.75 21h16.5M5.25 6h.008v.008H5.25V6zm0 3h.008v.008H5.25V9zm0 3h.008v.008H5.25V12zm4.5-6h.008v.008H9.75V6zm0 3h.008v.008H9.75V9zm0 3h.008v.008H9.75V12zm4.5-6h.008v.008h-.008V6zm0 3h.008v.008h-.008V9zm0 3h.008v.008h-.008V12zm4.5-6h.008v.008h-.008V6zm0 3h.008v.008h-.008V9zm0 3h.008v.008h-.008V12z"/></svg></span>
            <span>Clients</span>
        </a>
        <div class="sidebar-footer">
            <a href="{{ url_for('logout') }}" class="nav-item">
                <span class="nav-icon"><svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke-width="1.5" stroke="currentColor"><path stroke-linecap="round" stroke-linejoin="round" d="M15.75 9V5.25A2.25 2.25 0 0013.5 3h-6a2.25 2.25 0 00-2.25 2.25v13.5A2.25 2.25 0 007.5 21h6a2.25 2.25 0 002.25-2.25V15m3 0l3-3m0 0l-3-3m3 3H9"/></svg></span>
                <span>Logout</span>
            </a>
        </div>
    </nav>
    <main class="main-content">
        <div class="page-header">
            <h1>{{ page_title }}</h1>
        </div>
        <div id="toast-container" class="toast-container"></div>
        {% block content %}{% endblock %}
    </main>
    <script>
        function showToast(message, type = 'success') {
            const container = document.getElementById('toast-container');
            const toast = document.createElement('div');
            toast.className = 'toast ' + type;
            toast.textContent = message;
            container.appendChild(toast);
            setTimeout(() => toast.remove(), 4000);
        }
        function escapeHtml(text) {
            if (!text) return '';
            const div = document.createElement('div');
            div.textContent = text;
            return div.innerHTML;
        }
    </script>
    {% block scripts %}{% endblock %}
</body>
</html>
'''

LOGIN_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Login - Spike Admin</title>
    <style>
        :root {
            --bg-primary: #0f1117;
            --bg-card: #1e2130;
            --border-color: #2a2d3a;
            --text-primary: #e4e4e7;
            --text-muted: #71717a;
            --accent-primary: #6366f1;
            --accent-hover: #4f46e5;
            --accent-danger: #ef4444;
        }
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: var(--bg-primary);
            color: var(--text-primary);
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 100vh;
        }
        .login-card {
            background: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 40px;
            width: 100%;
            max-width: 400px;
        }
        .login-brand {
            text-align: center;
            font-size: 1.5rem;
            font-weight: 700;
            color: var(--accent-primary);
            margin-bottom: 32px;
        }
        .form-group { margin-bottom: 20px; }
        .form-group label {
            display: block;
            font-size: 0.85rem;
            color: var(--text-muted);
            margin-bottom: 8px;
        }
        .form-control {
            width: 100%;
            padding: 12px 16px;
            border-radius: 8px;
            border: 1px solid var(--border-color);
            background: var(--bg-primary);
            color: var(--text-primary);
            font-size: 0.95rem;
        }
        .form-control:focus { outline: none; border-color: var(--accent-primary); }
        .btn-login {
            width: 100%;
            padding: 12px;
            border-radius: 8px;
            border: none;
            background: var(--accent-primary);
            color: white;
            font-size: 1rem;
            font-weight: 600;
            cursor: pointer;
            transition: background 0.15s ease;
        }
        .btn-login:hover { background: var(--accent-hover); }
        .error-msg {
            background: rgba(239,68,68,0.1);
            border: 1px solid var(--accent-danger);
            border-radius: 8px;
            padding: 10px 14px;
            font-size: 0.85rem;
            color: var(--accent-danger);
            margin-bottom: 16px;
        }
    </style>
</head>
<body>
    <div class="login-card">
        <div class="login-brand">Spike Admin</div>
        {% if error %}
        <div class="error-msg">{{ error }}</div>
        {% endif %}
        <form method="POST">
            <div class="form-group">
                <label>Password</label>
                <input type="password" name="password" class="form-control" placeholder="Enter admin password" autofocus>
            </div>
            <button type="submit" class="btn-login">Login</button>
        </form>
    </div>
</body>
</html>
'''

DASHBOARD_TEMPLATE = '''
{% extends base %}
{% block content %}
<div class="stats-grid" id="statsGrid">
    <div class="stat-card">
        <div class="stat-value" id="totalSessions">-</div>
        <div class="stat-label">Total Interviews</div>
    </div>
    <div class="stat-card">
        <div class="stat-value" id="activeSessions">-</div>
        <div class="stat-label">Active</div>
    </div>
    <div class="stat-card">
        <div class="stat-value" id="totalQuestions">-</div>
        <div class="stat-label">Questions</div>
    </div>
</div>

<div class="card">
    <div class="card-header">
        <h2>Recent Interviews</h2>
    </div>
    <div class="table-container" id="recentInterviews">
        <div class="empty-state">
            <div class="empty-icon">⏳</div>
            <h3>Loading...</h3>
        </div>
    </div>
</div>
{% endblock %}

{% block scripts %}
<script>
document.addEventListener('DOMContentLoaded', loadDashboard);

async function loadDashboard() {
    try {
        const [sessionsRes, questionsRes] = await Promise.all([
            fetch('/api/sessions'),
            fetch('/api/questions')
        ]);
        const sessions = await sessionsRes.json();
        const questions = await questionsRes.json();

        document.getElementById('totalSessions').textContent = sessions.length;
        document.getElementById('activeSessions').textContent = sessions.length;
        document.getElementById('totalQuestions').textContent = questions.length;

        const container = document.getElementById('recentInterviews');
        if (sessions.length === 0) {
            container.innerHTML = '<div class="empty-state"><div class="empty-icon">🎤</div><h3>No interviews yet</h3></div>';
            return;
        }

        const recent = sessions.slice(0, 10);
        let html = '<table><thead><tr><th>Company</th><th>Target System</th><th>System Type</th><th>Date</th></tr></thead><tbody>';
        recent.forEach(s => {
            const dateStr = s.created_at ? new Date(s.created_at).toLocaleDateString() : 'N/A';
            const badgeClass = s.system_type ? s.system_type.toLowerCase().replace(' ', '-') : 'generic';
            html += '<tr>';
            html += '<td style="color:var(--text-primary);font-weight:500;">' + escapeHtml(s.company || 'N/A') + '</td>';
            html += '<td>' + escapeHtml(s.target_system || 'N/A') + '</td>';
            html += '<td><span class="badge badge-' + badgeClass + '">' + escapeHtml(s.system_type || 'N/A') + '</span></td>';
            html += '<td>' + dateStr + '</td>';
            html += '</tr>';
        });
        html += '</tbody></table>';
        container.innerHTML = html;
    } catch (err) {
        showToast('Failed to load dashboard: ' + err.message, 'error');
    }
}
</script>
{% endblock %}
'''

QUESTIONS_TEMPLATE = '''
{% extends base %}
{% block content %}
<div class="card">
    <div class="card-header">
        <h2>Interview Questions</h2>
        <button class="btn btn-primary" onclick="openCreateModal()">+ Add Question</button>
    </div>
    <div class="table-container" id="questionsTable">
        <div class="empty-state">
            <div class="empty-icon">⏳</div>
            <h3>Loading...</h3>
        </div>
    </div>
</div>

<div class="modal-overlay" id="questionModal">
    <div class="modal">
        <h2 id="modalTitle">Add Question</h2>
        <form id="questionForm" onsubmit="saveQuestion(event)">
            <input type="hidden" id="questionId">
            <div class="form-group">
                <label>Question Text</label>
                <textarea id="questionText" class="form-control" required></textarea>
            </div>
            <div class="form-group">
                <label>System Type</label>
                <select id="questionSystemType" class="form-control">
                    <option value="Generic">Generic</option>
                    <option value="AD-Azure">AD-Azure</option>
                    <option value="Target DB">Target DB</option>
                    <option value="SAP">SAP</option>
                    <option value="LDAP">LDAP</option>
                </select>
            </div>
            <div style="display:flex;gap:12px;justify-content:flex-end;margin-top:24px;">
                <button type="button" class="btn" style="background:var(--bg-hover);color:var(--text-secondary);" onclick="closeModal()">Cancel</button>
                <button type="submit" class="btn btn-primary">Save</button>
            </div>
        </form>
    </div>
</div>
{% endblock %}

{% block scripts %}
<script>
document.addEventListener('DOMContentLoaded', loadQuestions);

async function loadQuestions() {
    try {
        const response = await fetch('/api/questions');
        const questions = await response.json();
        const container = document.getElementById('questionsTable');

        if (questions.length === 0) {
            container.innerHTML = '<div class="empty-state"><div class="empty-icon">❓</div><h3>No questions yet</h3><p>Click "Add Question" to create your first interview question.</p></div>';
            return;
        }

        let html = '<table><thead><tr><th>ID</th><th>Question</th><th>System Type</th><th>Actions</th></tr></thead><tbody>';
        questions.forEach(q => {
            const badgeClass = q.system_type ? q.system_type.toLowerCase().replace(' ', '-') : 'generic';
            html += '<tr>';
            html += '<td>' + q.id + '</td>';
            html += '<td style="color:var(--text-primary);max-width:400px;">' + escapeHtml(q.question) + '</td>';
            html += '<td><span class="badge badge-' + badgeClass + '">' + escapeHtml(q.system_type) + '</span></td>';
            html += '<td><button class="btn btn-sm btn-primary" onclick="editQuestion(' + q.id + ')">Edit</button> ';
            html += '<button class="btn btn-sm btn-danger" onclick="deleteQuestion(' + q.id + ')">Delete</button></td>';
            html += '</tr>';
        });
        html += '</tbody></table>';
        container.innerHTML = html;
    } catch (err) {
        showToast('Failed to load questions: ' + err.message, 'error');
    }
}

function openCreateModal() {
    document.getElementById('modalTitle').textContent = 'Add Question';
    document.getElementById('questionId').value = '';
    document.getElementById('questionText').value = '';
    document.getElementById('questionSystemType').value = 'Generic';
    document.getElementById('questionModal').classList.add('active');
}

function closeModal() {
    document.getElementById('questionModal').classList.remove('active');
}

async function editQuestion(id) {
    try {
        const response = await fetch('/api/questions/' + id);
        const q = await response.json();
        document.getElementById('modalTitle').textContent = 'Edit Question';
        document.getElementById('questionId').value = q.id;
        document.getElementById('questionText').value = q.question;
        document.getElementById('questionSystemType').value = q.system_type;
        document.getElementById('questionModal').classList.add('active');
    } catch (err) {
        showToast('Failed to load question: ' + err.message, 'error');
    }
}

async function saveQuestion(e) {
    e.preventDefault();
    const id = document.getElementById('questionId').value;
    const data = {
        question: document.getElementById('questionText').value,
        system_type: document.getElementById('questionSystemType').value
    };

    try {
        const url = id ? '/api/questions/' + id : '/api/questions';
        const method = id ? 'PUT' : 'POST';
        const response = await fetch(url, {
            method: method,
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify(data)
        });
        if (!response.ok) throw new Error('Failed to save');
        closeModal();
        loadQuestions();
        showToast(id ? 'Question updated!' : 'Question created!', 'success');
    } catch (err) {
        showToast('Failed to save question: ' + err.message, 'error');
    }
}

async function deleteQuestion(id) {
    if (!confirm('Are you sure you want to delete this question?')) return;
    try {
        const response = await fetch('/api/questions/' + id, {method: 'DELETE'});
        if (!response.ok) throw new Error('Failed to delete');
        loadQuestions();
        showToast('Question deleted!', 'success');
    } catch (err) {
        showToast('Failed to delete question: ' + err.message, 'error');
    }
}
</script>
{% endblock %}
'''

INTERVIEWS_TEMPLATE = '''
{% extends base %}
{% block content %}
<div class="card">
    <div class="card-header">
        <h2>Interview Sessions</h2>
        <span style="font-size:0.85rem;color:var(--text-muted);" id="sessionCount"></span>
    </div>
    <div class="table-container" id="sessionsTable">
        <div class="empty-state">
            <div class="empty-icon">⏳</div>
            <h3>Loading...</h3>
        </div>
    </div>
</div>

<div class="modal-overlay" id="sessionModal">
    <div class="modal" style="max-width:700px;">
        <h2 id="sessionModalTitle">Interview Details</h2>
        <div id="sessionModalContent"></div>
        <div style="display:flex;justify-content:flex-end;margin-top:24px;">
            <button class="btn" style="background:var(--bg-hover);color:var(--text-secondary);" onclick="closeSessionModal()">Close</button>
        </div>
    </div>
</div>
{% endblock %}

{% block scripts %}
<script>
document.addEventListener('DOMContentLoaded', loadSessions);

async function loadSessions() {
    try {
        const response = await fetch('/api/sessions');
        const sessions = await response.json();
        const container = document.getElementById('sessionsTable');
        document.getElementById('sessionCount').textContent = sessions.length + ' total';

        if (sessions.length === 0) {
            container.innerHTML = '<div class="empty-state"><div class="empty-icon">🎤</div><h3>No interviews yet</h3><p>Interviews will appear here once users start the onboarding process.</p></div>';
            return;
        }

        let html = '<table><thead><tr><th>Company</th><th>Target System</th><th>System Type</th><th>Date</th><th>Actions</th></tr></thead><tbody>';
        sessions.forEach(s => {
            const dateStr = s.created_at ? new Date(s.created_at).toLocaleDateString() : 'N/A';
            const badgeClass = s.system_type ? s.system_type.toLowerCase().replace(' ', '-') : 'generic';
            html += '<tr>';
            html += '<td style="color:var(--text-primary);font-weight:500;">' + escapeHtml(s.company || 'N/A') + '</td>';
            html += '<td>' + escapeHtml(s.target_system || 'N/A') + '</td>';
            html += '<td><span class="badge badge-' + badgeClass + '">' + escapeHtml(s.system_type || 'N/A') + '</span></td>';
            html += '<td>' + dateStr + '</td>';
            html += '<td><button class="btn btn-sm btn-primary" onclick="viewSession(' + s.id + ')">View</button></td>';
            html += '</tr>';
        });
        html += '</tbody></table>';
        container.innerHTML = html;
    } catch (err) {
        showToast('Failed to load sessions: ' + err.message, 'error');
    }
}

async function viewSession(id) {
    try {
        const response = await fetch('/api/sessions/' + id);
        const s = await response.json();
        document.getElementById('sessionModalTitle').textContent = (s.company || 'Unknown') + ' - ' + (s.target_system || 'N/A');

        let html = '<div style="display:grid;gap:12px;">';
        html += '<div><strong>System Type:</strong> ' + escapeHtml(s.system_type || 'N/A') + '</div>';
        html += '<div><strong>Created:</strong> ' + (s.created_at ? new Date(s.created_at).toLocaleString() : 'N/A') + '</div>';

        if (s.collected_data_english && Object.keys(s.collected_data_english).length > 0) {
            html += '<div style="margin-top:16px;"><strong>Collected Data (English):</strong></div>';
            html += '<div style="background:var(--bg-primary);border-radius:8px;padding:16px;margin-top:8px;max-height:300px;overflow-y:auto;">';
            for (const [key, value] of Object.entries(s.collected_data_english)) {
                html += '<div style="margin-bottom:12px;"><div style="color:var(--accent-primary);font-size:0.8rem;font-weight:600;">' + escapeHtml(key) + '</div>';
                html += '<div style="color:var(--text-secondary);font-size:0.9rem;">' + escapeHtml(String(value)) + '</div></div>';
            }
            html += '</div>';
        } else if (s.collected_data_original && Object.keys(s.collected_data_original).length > 0) {
            html += '<div style="margin-top:16px;"><strong>Collected Data:</strong></div>';
            html += '<div style="background:var(--bg-primary);border-radius:8px;padding:16px;margin-top:8px;max-height:300px;overflow-y:auto;">';
            for (const [key, value] of Object.entries(s.collected_data_original)) {
                html += '<div style="margin-bottom:12px;"><div style="color:var(--accent-primary);font-size:0.8rem;font-weight:600;">' + escapeHtml(key) + '</div>';
                html += '<div style="color:var(--text-secondary);font-size:0.9rem;">' + escapeHtml(String(value)) + '</div></div>';
            }
            html += '</div>';
        }

        html += '</div>';
        document.getElementById('sessionModalContent').innerHTML = html;
        document.getElementById('sessionModal').classList.add('active');
    } catch (err) {
        showToast('Failed to load session: ' + err.message, 'error');
    }
}

function closeSessionModal() {
    document.getElementById('sessionModal').classList.remove('active');
}
</script>
{% endblock %}
'''

CLIENTS_TEMPLATE = '''
{% extends base %}
{% block content %}
<div class="card">
    <div class="card-header">
        <h2>Clients - Interview Overview</h2>
        <span style="font-size:0.85rem;color:var(--text-muted);" id="clientCount">Loading...</span>
    </div>
    <div id="clientsContainer">
        <div class="empty-state">
            <div class="empty-icon">⏳</div>
            <h3>Loading clients...</h3>
        </div>
    </div>
</div>
{% endblock %}

{% block scripts %}
<script>
document.addEventListener('DOMContentLoaded', loadClients);

async function loadClients() {
    try {
        const response = await fetch('/api/sessions');
        if (!response.ok) throw new Error('Failed to load sessions');
        const sessions = await response.json();

        const grouped = {};
        sessions.forEach(s => {
            const company = s.company || 'Unknown';
            if (!grouped[company]) grouped[company] = [];
            grouped[company].push(s);
        });

        const container = document.getElementById('clientsContainer');
        const countEl = document.getElementById('clientCount');
        const companies = Object.keys(grouped);
        countEl.textContent = companies.length + ' client' + (companies.length !== 1 ? 's' : '');

        if (companies.length === 0) {
            container.innerHTML = '<div class="empty-state"><div class="empty-icon">🏢</div><h3>No clients found</h3><p>Clients will appear here once interviews are completed.</p></div>';
            return;
        }

        let html = '';
        companies.forEach(company => {
            const companySessions = grouped[company];
            html += '<div class="card" style="margin-bottom:20px;">';
            html += '<div class="card-header"><h3>' + escapeHtml(company) + '</h3>';
            html += '<button class="btn btn-primary btn-sm" data-company="' + escapeHtml(company) + '" onclick="generateAllSolutionDesigns(this.dataset.company, this)">Generate All Solution Designs</button>';
            html += '</div>';
            html += '<div class="table-container"><table><thead><tr><th>Target System</th><th>System Type</th><th>Date</th><th>Action</th></tr></thead><tbody>';
            companySessions.forEach(s => {
                const dateStr = s.created_at ? new Date(s.created_at).toLocaleDateString('en-US', {year:'numeric',month:'short',day:'numeric'}) : 'N/A';
                const badgeClass = s.system_type ? s.system_type.toLowerCase().replace(' ', '-') : 'generic';
                html += '<tr>';
                html += '<td style="color:var(--text-primary);font-weight:500;">' + escapeHtml(s.target_system || 'N/A') + '</td>';
                html += '<td><span class="badge badge-' + badgeClass + '">' + escapeHtml(s.system_type || 'N/A') + '</span></td>';
                html += '<td>' + dateStr + '</td>';
                html += '<td><button class="btn btn-primary btn-sm" onclick="generateSolutionDesign(' + s.id + ', this)"><svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke-width="1.5" stroke="currentColor" width="14" height="14"><path stroke-linecap="round" stroke-linejoin="round" d="M9.594 3.94c.09-.542.56-.94 1.11-.94h2.593c.55 0 1.02.398 1.11.94l.213 1.281c.063.374.313.686.645.87.074.04.147.083.22.127.324.196.72.257 1.075.124l1.217-.456a1.125 1.125 0 011.37.49l1.296 2.247a1.125 1.125 0 01-.26 1.431l-1.003.827c-.293.24-.438.613-.431.992a6.759 6.759 0 010 .255c-.007.378.138.75.43.99l1.005.828c.424.35.534.954.26 1.43l-1.298 2.247a1.125 1.125 0 01-1.369.491l-1.217-.456c-.355-.133-.75-.072-1.076.124a6.57 6.57 0 01-.22.128c-.331.183-.581.495-.644.869l-.213 1.28c-.09.543-.56.941-1.11.941h-2.594c-.55 0-1.02-.398-1.11-.94l-.213-1.281c-.062-.374-.312-.686-.644-.87a6.52 6.52 0 01-.22-.127c-.325-.196-.72-.257-1.076-.124l-1.217.456a1.125 1.125 0 01-1.369-.49l-1.297-2.247a1.125 1.125 0 01.26-1.431l1.004-.827c.292-.24.437-.613.43-.992a6.932 6.932 0 010-.255c.007-.378-.138-.75-.43-.99l-1.004-.828a1.125 1.125 0 01-.26-1.43l1.297-2.247a1.125 1.125 0 011.37-.491l1.216.456c.356.133.751.072 1.076-.124.072-.044.146-.087.22-.128.332-.183.582-.495.644-.869l.214-1.281z"/><path stroke-linecap="round" stroke-linejoin="round" d="M15 12a3 3 0 11-6 0 3 3 0 016 0z"/></svg> Generate</button></td>';
                html += '</tr>';
            });
            html += '</tbody></table></div></div>';
        });

        container.innerHTML = html;
    } catch (err) {
        showToast('Failed to load clients: ' + err.message, 'error');
    }
}

async function generateSolutionDesign(sessionId, btn) {
    const originalText = btn.innerHTML;
    btn.innerHTML = '<svg width="14" height="14" viewBox="0 0 24 24" style="animation:spin 1s linear infinite;" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10" stroke-dasharray="31.4 31.4" stroke-linecap="round"/></svg> Generating...';
    btn.disabled = true;
    btn.style.opacity = '0.7';

    try {
        const response = await fetch('/api/generate-solution-design/' + sessionId, {method: 'POST'});
        if (!response.ok) {
            const data = await response.json();
            throw new Error(data.error || 'Generation failed');
        }

        const blob = await response.blob();
        const url = window.URL.createObjectURL(blob);
        const a = document.createElement('a');
        a.href = url;
        const disposition = response.headers.get('Content-Disposition');
        a.download = disposition ? disposition.split('filename=')[1].replace(/"/g, '') : 'solution_design.docx';
        document.body.appendChild(a);
        a.click();
        window.URL.revokeObjectURL(url);
        a.remove();

        showToast('Solution Design generated successfully!', 'success');
    } catch (err) {
        showToast('Failed to generate: ' + err.message, 'error');
    } finally {
        btn.innerHTML = originalText;
        btn.disabled = false;
        btn.style.opacity = '1';
    }
}

async function generateAllSolutionDesigns(company, btn) {
    const originalText = btn.innerHTML;
    btn.innerHTML = '<svg width="14" height="14" viewBox="0 0 24 24" style="animation:spin 1s linear infinite;" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10" stroke-dasharray="31.4 31.4" stroke-linecap="round"/></svg> Generating All...';
    btn.disabled = true;
    btn.style.opacity = '0.7';

    try {
        const response = await fetch('/api/generate-all-solution-designs', {
            method: 'POST',
            headers: {'Content-Type': 'application/json'},
            body: JSON.stringify({company: company})
        });
        if (!response.ok) {
            const data = await response.json();
            throw new Error(data.error || 'Generation failed');
        }

        const blob = await response.blob();
        const url = window.URL.createObjectURL(blob);
        const a = document.createElement('a');
        a.href = url;
        const disposition = response.headers.get('Content-Disposition');
        a.download = disposition ? disposition.split('filename=')[1].replace(/"/g, '') : 'solution_designs_all.docx';
        document.body.appendChild(a);
        a.click();
        window.URL.revokeObjectURL(url);
        a.remove();

        showToast('All Solution Designs generated successfully!', 'success');
    } catch (err) {
        showToast('Failed to generate: ' + err.message, 'error');
    } finally {
        btn.innerHTML = originalText;
        btn.disabled = false;
        btn.style.opacity = '1';
    }
}
</script>
<style>
@keyframes spin { from { transform: rotate(0deg); } to { transform: rotate(360deg); } }
</style>
{% endblock %}'''

def render_page(template_str, **kwargs):
    """Render a child template that uses extends base and block tags."""
    content_match = re.search(r'\{% block content %\}(.*?)\{% endblock %\}', template_str, re.DOTALL)
    scripts_match = re.search(r'\{% block scripts %\}(.*?)\{% endblock %\}', template_str, re.DOTALL)

    content_block = content_match.group(1) if content_match else ''
    scripts_block = scripts_match.group(1) if scripts_match else ''

    full_template = BASE_TEMPLATE.replace('{% block content %}{% endblock %}', content_block)
    full_template = full_template.replace('{% block scripts %}{% endblock %}', scripts_block)

    return render_template_string(full_template, **kwargs)

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        password = request.form.get('password', '')
        if password == ADMIN_PASSWORD:
            session['authenticated'] = True
            return redirect(url_for('dashboard'))
        else:
            error = 'Invalid password. Please try again.'
    return render_template_string(LOGIN_TEMPLATE, error=error)

@app.route('/logout')
def logout():
    session.pop('authenticated', None)
    return redirect(url_for('login'))

@app.route('/')
@login_required
def dashboard():
    return render_page(
        DASHBOARD_TEMPLATE,
        page_title='Dashboard',
        active_page='dashboard',
    )

@app.route('/questions')
@login_required
def questions_page():
    return render_page(
        QUESTIONS_TEMPLATE,
        page_title='Questions',
        active_page='questions',
    )

@app.route('/interviews')
@login_required
def interviews_page():
    return render_page(
        INTERVIEWS_TEMPLATE,
        page_title='Interviews',
        active_page='interviews',
    )

@app.route('/clients')
@login_required
def clients_page():
    return render_page(
        CLIENTS_TEMPLATE,
        page_title='Clients',
        active_page='clients',
    )

@app.route('/api/questions', methods=['GET'])
@login_required
def api_get_questions():
    with get_db() as db:
        questions = db.query(Question).order_by(Question.id.asc()).all()
        return jsonify([{
            'id': q.id,
            'question': q.question,
            'system_type': q.system_type
        } for q in questions])

@app.route('/api/questions', methods=['POST'])
@login_required
def api_create_question():
    data = request.get_json()
    if not data or not data.get('question'):
        return jsonify({'error': 'question is required'}), 400

    with get_db() as db:
        question = Question(
            question=data['question'],
            system_type=data.get('system_type', 'Generic')
        )
        db.add(question)
        db.commit()
        db.refresh(question)
        return jsonify({
            'id': question.id,
            'question': question.question,
            'system_type': question.system_type
        }), 201

@app.route('/api/questions/<int:question_id>', methods=['GET'])
@login_required
def api_get_question(question_id):
    with get_db() as db:
        q = db.query(Question).filter(Question.id == question_id).first()
        if not q:
            return jsonify({'error': 'Question not found'}), 404
        return jsonify({
            'id': q.id,
            'question': q.question,
            'system_type': q.system_type
        })

@app.route('/api/questions/<int:question_id>', methods=['PUT'])
@login_required
def api_update_question(question_id):
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    with get_db() as db:
        q = db.query(Question).filter(Question.id == question_id).first()
        if not q:
            return jsonify({'error': 'Question not found'}), 404

        if 'question' in data:
            q.question = data['question']
        if 'system_type' in data:
            q.system_type = data['system_type']

        db.commit()
        db.refresh(q)
        return jsonify({
            'id': q.id,
            'question': q.question,
            'system_type': q.system_type
        })

@app.route('/api/questions/<int:question_id>', methods=['DELETE'])
@login_required
def api_delete_question(question_id):
    with get_db() as db:
        q = db.query(Question).filter(Question.id == question_id).first()
        if not q:
            return jsonify({'error': 'Question not found'}), 404
        db.delete(q)
        db.commit()
        return jsonify({'message': 'Question deleted'})

@app.route('/api/sessions', methods=['GET'])
@login_required
def api_get_sessions():
    with get_db() as db:
        sessions = db.query(OnboardingSession).order_by(OnboardingSession.created_at.desc()).all()
        result = []
        for s in sessions:
            collected_original = s.collected_data_original
            collected_english = s.collected_data_english
            if isinstance(collected_original, str):
                try:
                    collected_original = json.loads(collected_original)
                except:
                    collected_original = {}
            if isinstance(collected_english, str):
                try:
                    collected_english = json.loads(collected_english)
                except:
                    collected_english = {}

            result.append({
                'id': s.id,
                'company': s.company,
                'target_system': s.target_system,
                'system_type': s.system_type,
                'created_at': s.created_at.isoformat() if s.created_at else None,
                'collected_data_original': collected_original or {},
                'collected_data_english': collected_english or {}
            })
        return jsonify(result)

@app.route('/api/sessions/<int:session_id>', methods=['GET'])
@login_required
def api_get_session(session_id):
    with get_db() as db:
        s = db.query(OnboardingSession).filter(OnboardingSession.id == session_id).first()
        if not s:
            return jsonify({'error': 'Session not found'}), 404

        collected_original = s.collected_data_original
        collected_english = s.collected_data_english
        if isinstance(collected_original, str):
            try:
                collected_original = json.loads(collected_original)
            except:
                collected_original = {}
        if isinstance(collected_english, str):
            try:
                collected_english = json.loads(collected_english)
            except:
                collected_english = {}

        return jsonify({
            'id': s.id,
            'company': s.company,
            'target_system': s.target_system,
            'system_type': s.system_type,
            'created_at': s.created_at.isoformat() if s.created_at else None,
            'collected_data_original': collected_original or {},
            'collected_data_english': collected_english or {}
        })

def _build_llm_prompt(session_obj):
    """Build a strict, evidence-grounded LLM prompt from one interview session."""
    original = session_obj.collected_data_original
    english = session_obj.collected_data_english

    if isinstance(original, str):
        try:
            original = json.loads(original)
        except Exception:
            original = {}

    if isinstance(english, str):
        try:
            english = json.loads(english)
        except Exception:
            english = {}

    interview_data = english or original or {}
    interview_data = {
        str(k).strip(): str(v).strip()
        for k, v in interview_data.items()
        if str(k).strip() and str(v).strip()
    }

    company = session_obj.company or "[TO BE DEFINED]"
    target_system = session_obj.target_system or "[TO BE DEFINED]"
    system_type = session_obj.system_type or "[TO BE DEFINED]"

    qa_text = "\n".join(
        f"{idx}. Question: {question}\n   Answer: {answer}"
        for idx, (question, answer) in enumerate(interview_data.items(), start=1)
    ) or "[TO BE DEFINED]"

    template_filled = SOLUTION_DESIGN_TEMPLATE_PROMPT.format(
        company=company,
        target_system=target_system,
        system_type=system_type,
    )

    if system_type == "AD-Azure":
        system_mapping = """
SYSTEM-SPECIFIC MAPPING FOR AD-Azure:
Use interview answers about AD domains, OU hierarchy, employee/external OUs, disabled-account OUs, group OUs, group descriptions, standard/default groups, Office 365 license assignment, Azure AD Connect schedule, synchronization scope, network/firewall constraints, sample employee/external accounts, SAMAccountName and UPN conventions, account creation/modification/disabling, OU movement rules, exceptions, and custom scenarios.

Map them mainly to:
- Current State
- Project Environments
- Integration with Target Systems
- Connector and Connection Configuration
- Account Definition and Identity Linking
- Attribute Mapping Logic
- Access Rights / Entitlements
- Synchronization and Provisioning
- RBAC Model
- Joiner, Mover and Leaver Processes
"""
    if system_type == "ldap":
        system_mapping = """
SYSTEM-SPECIFIC MAPPING FOR LDAP:
Use interview answers about LDAP domains, DN hierarchy, employee/external DNs, disabled-account DNs, group DNs, group descriptions, standard/default groups, synchronization scope, network/firewall constraints, sample employee/external accounts, DN and names conventions, account creation/modification/disabling, DN movement rules, exceptions, and custom scenarios.

Map them mainly to:
- Current State
- Project Environments
- Integration with Target Systems
- Connector and Connection Configuration
- Account Definition and Identity Linking
- Attribute Mapping Logic
- Access Rights / Entitlements
- Synchronization and Provisioning
- RBAC Model
- Joiner, Mover and Leaver Processes
"""
    elif system_type == "Target DB":
        system_mapping = """
SYSTEM-SPECIFIC MAPPING FOR Target DB:
Use interview answers about deployment model, network location, authentication mechanism, authorization mechanism, database technology, country/business-unit instances, ongoing projects/upgrades, local user records, unique user key, multiple account/profile behavior, required AD groups for access, user profiling model, access-right tables, access-right attributes, user-friendly descriptions, assignment tables, employees/external users, departments with access, departments excluded from access, first-access process, account modification process, disabled/terminated user behavior, SoD rules, custom scenarios, connector service account, environments, mandatory/optional user attributes, user status attribute and values, and access-right naming convention.

Map them mainly to:
- Current State
- Project Environments
- Integration with Target Systems
- Connector and Connection Configuration
- Account Definition and Identity Linking
- Attribute Mapping Logic
- Access Rights / Entitlements
- Synchronization and Provisioning
- RBAC Model
- Joiner, Mover and Leaver Processes

Important Target DB rule:
If the interview mentions an Active Directory group required for basic access, that group must be explicitly included under Access Rights / Entitlements and, if relevant, under Synchronization and Provisioning.
"""
    elif system_type == "SAP":
        system_mapping = """
SYSTEM-SPECIFIC MAPPING FOR SAP:
Use interview answers about SAP version, SAP type, authentication, authorization, roles, profiles, modules, environments, connectors, technical users, provisioning, SoD, and custom scenarios.
If a detail is not available, use [TO BE DEFINED].
"""
    else:
        system_mapping = """
SYSTEM-SPECIFIC MAPPING FOR Generic:
Infer the integration approach only from explicit interview answers.
If connector type or integration mode cannot be inferred, use [TO BE DEFINED].
"""

    prompt = f"""
You are a Senior IAM / IGA Solution Architect and Senior Consultant.

Your task is to generate a professional Solution Design document based on one technical onboarding interview.

You must behave like a strict enterprise IAM consultant:
- Extract every relevant piece of information from the interview.
- Map each interview answer to the correct Solution Design section.
- Never omit information that is present in the interview.
- Never invent customer-specific information.
- Use [TO BE DEFINED] only when the information is not available in the interview.
- If information is partially available, write the available part and use [TO BE DEFINED] only for the missing sub-fields.
- Do not use [TO BE DEFINED] as a generic fallback if the interview provides usable information.
- Do not generate assumptions, fake details, fake server names, fake owners, fake schedules, fake workflows, fake architectures, fake authoritative sources, fake privileges, fake tables, fake fields, fake naming conventions, or fake processes.

SOURCE OF TRUTH

The only source of customer-specific truth is the interview data provided below.

You may use the Solution Design template only for:
- structure;
- standard section headings;
- generic One Identity Manager / IGA explanatory text;
- document style and professional tone.

You must not use the template to invent customer-specific facts.

SESSION METADATA

Company:
{company}

Target System:
{target_system}

System Type:
{system_type}

INTERVIEW DATA

{qa_text}

SOLUTION DESIGN TEMPLATE / STRUCTURE TO FOLLOW

{template_filled}

SYSTEM-TYPE SPECIFIC MAPPING RULES

{system_mapping}

MANDATORY ANALYSIS PROCESS

Before writing the final Solution Design, internally perform these steps:

1. Read all interview questions and answers carefully.
2. Extract all factual items from the interview.
   Examples of factual items:
   - system deployment model;
   - network location;
   - authentication mechanism;
   - authorization mechanism;
   - database technology;
   - environments;
   - instances;
   - tables;
   - fields;
   - keys;
   - usernames;
   - groups;
   - roles;
   - permissions;
   - service accounts;
   - privileges;
   - synchronization schedules;
   - lifecycle processes;
   - provisioning processes;
   - deactivation rules;
   - SoD constraints;
   - exceptions;
   - ongoing projects;
   - upgrade plans;
   - naming conventions;
   - owners;
   - approval workflows.

3. For every factual item, decide where it belongs in the Solution Design.
4. Ensure that every non-empty interview answer is represented somewhere in the final document.
5. If an answer is relevant to multiple sections, include it in all necessary sections without creating contradictions.
6. Only after this internal mapping, generate the Solution Design.

CRITICAL COMPLETENESS RULE

Every answer in the interview must be used.

If the interview contains a specific value such as:
- an AD group name;
- a database table;
- a database field;
- a service account;
- a role/profile name;
- a permission;
- an environment;
- a database version;
- a network zone;
- a synchronization frequency;
- a lifecycle rule;
- an exception;
- an upgrade or ongoing project;

then that value must appear explicitly in the final Solution Design.

Do not summarize away important technical details.

For example:
- If the interview mentions an AD group required for basic access, include it under Access Rights / Entitlements and, if relevant, under Synchronization and Provisioning.
- If the interview mentions a service account and its permissions, include both the account name and the exact permissions.
- If the interview mentions a table used for user records, role assignments, or access rights, include the exact table name.
- If the interview mentions a unique user key, include it under Account Definition and Identity Linking.
- If the interview mentions environments, include all of them under Project Environments.
- If the interview mentions an upgrade or parallel project, include it under Current State, Constraints, Risks, or Integration Considerations.

STRICT NON-HALLUCINATION RULE

You must not create any customer-specific detail that is not supported by the interview.

Incorrect behavior:
- Writing a synchronization schedule when none was provided.
- Inventing a server name, schema name, DB host, application owner, approval workflow, lifecycle timeline, or account definition name.
- Assuming that a process is automatic when the interview does not say so.
- Assuming an authoritative source if the interview does not mention one.
- Assuming SoD rules if the interview does not describe them.
- Assuming deactivation or deletion timing if not provided.

Correct behavior:
- Write the known details.
- For the missing part, write exactly: [TO BE DEFINED]

PARTIAL INFORMATION RULE

If a section asks for multiple details and the interview provides only some of them, do not mark the full section as [TO BE DEFINED].

Instead, write the available details and mark only the missing ones.

Example:

Connector and Connection Configuration:
- Connector type: Generic Database Connector
- Database technology: Oracle 19c Enterprise Edition
- Service account: SVC_IGA_CONNECTOR
- Expected privileges: SELECT, INSERT, UPDATE, DELETE on USERS and USER_ROLES
- Server name: [TO BE DEFINED]
- Schema name: [TO BE DEFINED]
- Synchronization schedule: [TO BE DEFINED]

OUTPUT REQUIREMENTS

Generate only the final Solution Design in markdown format.

Do not output:
- your internal analysis;
- the evidence mapping table;
- reasoning steps;
- comments about the prompt;
- code fences;
- JSON;
- explanations outside the document.

The document must be professional, structured, and suitable for an enterprise IAM / IGA Solution Design.

Use the same section order as the provided Solution Design template.

FINAL INTERNAL QUALITY CHECK BEFORE ANSWERING

Before producing the final answer, internally verify:

1. Every interview answer has been used at least once.
2. No customer-specific information has been invented.
3. Every missing field is marked exactly as [TO BE DEFINED].
4. Partial information has not been discarded.
5. Exact technical names from the interview are preserved.
6. No section contradicts the interview.
7. Generic IGA text is clearly separated from customer-specific facts.
8. All target-system-specific details are placed in the appropriate sections.

If any interview answer is not represented, revise the document before finalizing.
"""
    return prompt

def _call_llm(prompt):
    """Call Azure OpenAI API and return the generated markdown content."""
    if not AZURE_API_KEY:
        raise Exception("Missing AZURE_API_KEY environment variable. Set it in .env file.")

    headers = {
        "Authorization": f"Bearer {AZURE_API_KEY}",
        "Content-Type": "application/json",
    }

    system_message = (
        "You are an expert IAM Solution Architect. "
        "Generate professional Solution Design documents in markdown. "
        "You must ground every populated detail in the provided interview data. "
        "Before using [TO BE DEFINED], verify that the interview does not provide that information. "
        "Never invent or assume customer-specific data."
    )

    payload = {
        "input": [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ],
        "model": AZURE_LLM_MODEL,
    }

    logger.info("Calling Azure LLM for Solution Design generation")
    response = requests.post(
        AZURE_LLM_URL,
        headers=headers,
        json=payload,
        verify=False,
        timeout=180,
    )
    logger.info("Azure LLM response status: %s", response.status_code)
    response.raise_for_status()

    result = response.json()
    return result["output"][0]["content"][0]["text"]

def _add_formatted_paragraph(doc, text):
    """Add a paragraph with bold formatting support (**text**)."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    paragraph = doc.add_paragraph()
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
    return paragraph

def _add_content_to_doc(doc, content):
    """Parse markdown-like LLM content and add to a python-docx Document."""
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph('')
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # Table detection
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            # Collect all table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Filter out separator lines (|---|---|)
            data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l)]

            if data_lines:
                # Parse cells
                rows = []
                for tl in data_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    rows.append(cells)

                if rows:
                    num_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                table.rows[row_idx].cells[col_idx].text = cell_text

                    # Bold first row (header)
                    if rows:
                        for cell in table.rows[0].cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

                    doc.add_paragraph('')
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Regular paragraph (with bold support)
        if '**' in stripped:
            _add_formatted_paragraph(doc, stripped)
        else:
            doc.add_paragraph(stripped)

        i += 1

@app.route('/api/generate-solution-design/<int:session_id>', methods=['POST'])
@login_required
def api_generate_solution_design(session_id):
    try:
        with get_db() as db:
            s = db.query(OnboardingSession).filter(OnboardingSession.id == session_id).first()
            if not s:
                return jsonify({'error': 'Session not found'}), 404

            prompt = _build_llm_prompt(s)
            content = _call_llm(prompt)

            # Generate Word document
            doc = Document()

            # Title page
            title = doc.add_heading(f'{s.company or "Company"} - IGA - Solution Design', level=0)
            doc.add_paragraph('')
            doc.add_paragraph(f'Target System: {s.target_system or "N/A"}')
            doc.add_paragraph(f'System Type: {s.system_type or "N/A"}')
            doc.add_paragraph('')

            # Version table
            version_table = doc.add_table(rows=2, cols=4)
            version_table.style = 'Table Grid'
            headers = ['VERSION', 'CHANGES', 'AUTHOR', 'DATE']
            for idx, header in enumerate(headers):
                version_table.rows[0].cells[idx].text = header
                for paragraph in version_table.rows[0].cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            version_table.rows[1].cells[0].text = '1.0'
            version_table.rows[1].cells[1].text = 'First draft of the document'
            version_table.rows[1].cells[2].text = '[TO BE DEFINED]'
            version_table.rows[1].cells[3].text = datetime.now().strftime("%d/%m/%Y")

            doc.add_page_break()

            # Add LLM-generated content
            _add_content_to_doc(doc, content)

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Generate filename
            company_clean = re.sub(r'[^a-zA-Z0-9]', '_', s.company or 'unknown')
            target_clean = re.sub(r'[^a-zA-Z0-9]', '_', s.target_system or 'system')
            filename = f"{company_clean}_IGA_Solution_Design_{target_clean}.docx"

            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    except requests.exceptions.RequestException as e:
        return jsonify({'error': f'LLM service error: {str(e)}'}), 502
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-all-solution-designs', methods=['POST'])
@login_required
def api_generate_all_solution_designs():
    try:
        data = request.get_json()
        if not data or not data.get('company'):
            return jsonify({'error': 'company is required'}), 400

        company = data['company']

        with get_db() as db:
            sessions = db.query(OnboardingSession).filter(
                OnboardingSession.company == company
            ).order_by(OnboardingSession.created_at.desc()).all()

            if not sessions:
                return jsonify({'error': 'No sessions found for this company'}), 404

            # Generate combined Word document
            doc = Document()

            # Title page
            doc.add_heading(f'{company} - IGA - Solution Design', level=0)
            doc.add_paragraph('')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total Target Systems: {len(sessions)}')
            doc.add_paragraph('')

            # Version table
            version_table = doc.add_table(rows=2, cols=4)
            version_table.style = 'Table Grid'
            headers = ['VERSION', 'CHANGES', 'AUTHOR', 'DATE']
            for idx, header in enumerate(headers):
                version_table.rows[0].cells[idx].text = header
                for paragraph in version_table.rows[0].cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            version_table.rows[1].cells[0].text = '1.0'
            version_table.rows[1].cells[1].text = 'First draft of the document'
            version_table.rows[1].cells[2].text = '[TO BE DEFINED]'
            version_table.rows[1].cells[3].text = datetime.now().strftime("%d/%m/%Y")

            # Table of contents placeholder
            doc.add_page_break()
            doc.add_heading('Table of Contents', level=1)
            for idx, s in enumerate(sessions):
                doc.add_paragraph(
                    f'{idx + 1}. {s.target_system or "N/A"} ({s.system_type or "N/A"})',
                    style='List Number'
                )

            # Generate content for each session
            for idx, s in enumerate(sessions):
                doc.add_page_break()
                doc.add_heading(
                    f'{s.target_system or "N/A"} ({s.system_type or "N/A"})',
                    level=1
                )
                doc.add_paragraph(f'Target System: {s.target_system or "N/A"}')
                doc.add_paragraph(f'System Type: {s.system_type or "N/A"}')
                doc.add_paragraph('')

                prompt = _build_llm_prompt(s)
                content = _call_llm(prompt)
                _add_content_to_doc(doc, content)

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Generate filename
            company_clean = re.sub(r'[^a-zA-Z0-9]', '_', company)
            filename = f"{company_clean}_IGA_Solution_Design_All_Systems.docx"

            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    except requests.exceptions.RequestException as e:
        return jsonify({'error': f'LLM service error: {str(e)}'}), 502
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': 'Internal server error'}), 500
